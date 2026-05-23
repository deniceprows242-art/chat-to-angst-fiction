#!/usr/bin/env python3
"""Convert a markdown file to DOCX with Chinese font support.

Usage:
    python generate_docx.py <input.md> [output.docx]

If output is not specified, it replaces .md with .docx.
"""

import os
import re
import sys

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


FONT_NAME = "微软雅黑"
FONT_SIZE = Pt(10.5)
LINE_SPACING = 1.5


def add_paragraph(doc, text):
    """Add a paragraph, handling **bold** and ***bold-italic*** segments."""
    if not text.strip():
        doc.add_paragraph("")
        return

    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.line_spacing = LINE_SPACING

    segments = re.split(r"(\*\*\*.*?\*\*\*|\*\*.*?\*\*)", text)
    for segment in segments:
        run = paragraph.add_run("")
        if segment.startswith("***") and segment.endswith("***"):
            run.text = segment[3:-3]
            run.bold = True
            run.italic = True
        elif segment.startswith("**") and segment.endswith("**"):
            run.text = segment[2:-2]
            run.bold = True
        elif segment == "---":
            run.text = "-" * 20
        else:
            run.text = segment
        run.font.name = FONT_NAME
        run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
        run.font.size = FONT_SIZE


def convert(input_path, output_path=None):
    if output_path is None:
        base, _ = os.path.splitext(input_path)
        output_path = base + ".docx"

    if not os.path.exists(input_path):
        print(f"Error: {input_path} not found.")
        sys.exit(1)

    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = FONT_NAME
    style.font.size = FONT_SIZE
    style.element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
    style.paragraph_format.line_spacing = LINE_SPACING

    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    with open(input_path, "r", encoding="utf-8") as file:
        for line in file:
            add_paragraph(doc, line.rstrip("\n").rstrip("\r"))

    doc.save(output_path)
    print(f"Saved: {output_path}")
    return output_path


if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Usage: python generate_docx.py <input.md> [output.docx]")
        sys.exit(1)
    convert(sys.argv[1], sys.argv[2] if len(sys.argv) > 2 else None)
